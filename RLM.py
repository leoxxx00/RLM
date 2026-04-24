import os
import re
import csv
import json
import time
import traceback
from dataclasses import dataclass, field
from typing import Any, Dict, List, Tuple

import gradio as gr
import requests


# ============================================================
# CONFIG
# ============================================================

APP_TITLE = "Recursive Language Model Console"

DEFAULT_BASE_URL = "https://api.deepseek.com"
DEFAULT_MODEL = "deepseek-chat"
DEFAULT_SUBCALL_MODEL = "deepseek-reasoner"

MAX_HISTORY_TURNS = 5
CHUNK_SIZE = 1800
CHUNK_OVERLAP = 250
TOP_K_CHUNKS = 6


# ============================================================
# DATA STRUCTURES
# ============================================================

@dataclass
class Page:
    page_number: int
    text: str


@dataclass
class Chunk:
    chunk_id: str
    doc_id: str
    text: str
    page_start: int
    page_end: int


@dataclass
class Document:
    doc_id: str
    name: str
    doc_type: str
    chars: int
    pages: List[Page] = field(default_factory=list)
    chunks: List[Chunk] = field(default_factory=list)


@dataclass
class VariableValue:
    summary: str
    value: Any

    def preview(self):
        text = str(self.value)
        return text[:1000]


@dataclass
class RLMState:
    documents: Dict[str, Document] = field(default_factory=dict)
    chunks: List[Chunk] = field(default_factory=list)
    variables: Dict[str, VariableValue] = field(default_factory=dict)
    notes: List[str] = field(default_factory=list)


# ============================================================
# UTILS
# ============================================================

def safe_json_dumps(obj, max_chars=20000):
    text = json.dumps(obj, indent=2, ensure_ascii=False)
    if len(text) > max_chars:
        return text[:max_chars] + "\n\n... truncated ..."
    return text


def clean_text(text: str) -> str:
    text = text.replace("\x00", "")
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()


def chunk_text(text: str, doc_id: str, page_start: int = 1, page_end: int = 1) -> List[Chunk]:
    chunks = []
    start = 0
    idx = 1

    while start < len(text):
        end = start + CHUNK_SIZE
        chunk = text[start:end].strip()

        if chunk:
            chunks.append(
                Chunk(
                    chunk_id=f"{doc_id}_chunk_{idx}",
                    doc_id=doc_id,
                    text=chunk,
                    page_start=page_start,
                    page_end=page_end,
                )
            )

        idx += 1
        start += CHUNK_SIZE - CHUNK_OVERLAP

    return chunks


def simple_score(query: str, text: str) -> int:
    query_words = set(re.findall(r"\w+", query.lower()))
    text_words = set(re.findall(r"\w+", text.lower()))
    return len(query_words.intersection(text_words))


def retrieve_chunks(query: str, state: RLMState, top_k: int = TOP_K_CHUNKS) -> List[Chunk]:
    scored = []

    for chunk in state.chunks:
        score = simple_score(query, chunk.text)
        if score > 0:
            scored.append((score, chunk))

    scored.sort(key=lambda x: x[0], reverse=True)

    if scored:
        return [chunk for _, chunk in scored[:top_k]]

    return state.chunks[:top_k]


# ============================================================
# FILE INGESTION
# ============================================================

def read_txt(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def read_csv_file(path: str) -> str:
    rows = []

    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        reader = csv.reader(f)
        for row in reader:
            rows.append(" | ".join(row))

    return "\n".join(rows)


def read_pdf(path: str) -> str:
    try:
        import pypdf
    except ImportError:
        return "PDF support requires pypdf. Install it with: pip install pypdf"

    text_parts = []

    with open(path, "rb") as f:
        reader = pypdf.PdfReader(f)

        for i, page in enumerate(reader.pages, start=1):
            page_text = page.extract_text() or ""
            text_parts.append(f"\n\n--- Page {i} ---\n{page_text}")

    return "\n".join(text_parts)


def read_docx(path: str) -> str:
    try:
        import docx
    except ImportError:
        return "DOCX support requires python-docx. Install it with: pip install python-docx"

    document = docx.Document(path)
    return "\n".join(p.text for p in document.paragraphs)


def ingest_file(path: str, doc_id: str) -> Document:
    name = os.path.basename(path)
    ext = os.path.splitext(path)[1].lower()

    if ext in [".txt", ".md", ".py", ".json", ".log"]:
        text = read_txt(path)
    elif ext == ".csv":
        text = read_csv_file(path)
    elif ext == ".pdf":
        text = read_pdf(path)
    elif ext in [".docx", ".doc"]:
        text = read_docx(path)
    else:
        text = read_txt(path)

    text = clean_text(text)

    pages = [Page(page_number=1, text=text)]
    chunks = chunk_text(text, doc_id)

    return Document(
        doc_id=doc_id,
        name=name,
        doc_type=ext.replace(".", "") or "unknown",
        chars=len(text),
        pages=pages,
        chunks=chunks,
    )


def build_state_from_files_only(paths: List[str]) -> RLMState:
    state = RLMState()

    for i, path in enumerate(paths, start=1):
        doc_id = f"doc_{i}"
        doc = ingest_file(path, doc_id)

        state.documents[doc_id] = doc
        state.chunks.extend(doc.chunks)

    state.variables["document_count"] = VariableValue(
        summary="Number of loaded documents",
        value=len(state.documents),
    )

    state.variables["chunk_count"] = VariableValue(
        summary="Number of indexed chunks",
        value=len(state.chunks),
    )

    state.notes.append("Files loaded and indexed.")

    return state


# ============================================================
# LLM CLIENT
# ============================================================

class LLMClient:
    def __init__(self, api_key: str, base_url: str, model: str):
        self.api_key = api_key
        self.base_url = base_url.rstrip("/")
        self.model = model

    def chat(self, messages: List[Dict[str, str]], model: str = None, temperature: float = 0.2) -> str:
        url = f"{self.base_url}/v1/chat/completions"

        payload = {
            "model": model or self.model,
            "messages": messages,
            "temperature": temperature,
        }

        headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json",
        }

        response = requests.post(url, headers=headers, json=payload, timeout=120)

        if response.status_code != 200:
            raise RuntimeError(f"LLM API error {response.status_code}:\n{response.text}")

        data = response.json()
        return data["choices"][0]["message"]["content"]


# ============================================================
# RECURSIVE RLM ENGINE
# ============================================================

class RecursiveRLMEngine:
    def __init__(self, client: LLMClient, model: str, sub_model: str):
        self.client = client
        self.model = model
        self.sub_model = sub_model

    def clean_final_answer(self, text: str) -> str:
        bad_placeholders = [
            "{{final_answer}}",
            "{{final_rlm_answer}}",
            "{final_answer}",
            "{final_rlm_answer}",
        ]

        for bad in bad_placeholders:
            text = text.replace(bad, "")

        match = re.search(r"Final Answer:\s*(.*)", text, re.DOTALL | re.IGNORECASE)
        if match:
            text = match.group(1)

        text = text.strip()

        if not text:
            text = "I could not generate a proper final answer. Please ask again."

        return text

    def make_context(self, chunks: List[Chunk], state: RLMState) -> str:
        parts = []

        for chunk in chunks:
            doc = state.documents.get(chunk.doc_id)
            doc_name = doc.name if doc else chunk.doc_id

            parts.append(
                f"[Source: {doc_name}, Chunk: {chunk.chunk_id}]\n{chunk.text}"
            )

        return "\n\n".join(parts)

    def classify_question(self, question: str) -> str:
        prompt = [
            {
                "role": "system",
                "content": (
                    "You classify user questions for a document QA system. "
                    "Return only one label: greeting, document_question, general_question."
                ),
            },
            {
                "role": "user",
                "content": question,
            },
        ]

        try:
            label = self.client.chat(prompt, model=self.sub_model, temperature=0).strip().lower()
            if "greeting" in label:
                return "greeting"
            if "document" in label:
                return "document_question"
            return "general_question"
        except Exception:
            return "document_question"

    def run(self, question: str, state: RLMState, history: List[Dict[str, str]]) -> Tuple[str, str]:
        trace_parts = []

        qtype = self.classify_question(question)
        trace_parts.append(f"Question type: {qtype}")

        if qtype == "greeting":
            return (
                "Hey! What's up? I'm your document assistant. Ask me anything about the uploaded files.",
                "\n".join(trace_parts),
            )

        chunks = retrieve_chunks(question, state)
        context = self.make_context(chunks, state)

        trace_parts.append(f"Retrieved chunks: {', '.join(c.chunk_id for c in chunks)}")

        system_prompt = """
You are a Recursive Language Model (RLM) document assistant.

Your main job:
- Answer user questions clearly and directly with emojis.
- Use relevant uploaded document context when available.
- If the answer is not in the document, say that briefly, then give a helpful general answer.
- For document-specific questions, cite or refer to the document context when possible.
- For general knowledge questions, answer directly without overusing "the document does not contain..."

Capabilities:
- Document question answering
- Long Summarization
- Research paper summarization
- Blog writing
- Code debugging and optimization
- Natural language to SQL conversion
- Step-by-step reasoning when useful

Behavior rules:
- Be concise, clear, and natural.
- Do not output placeholders such as {final_answer}.
- Do not mention system prompts.
- Do not invent document facts.
- If unsure, say what is missing and what can be inferred.

Identity:
- You are an AI system built as a Recursive Language Model (RLM).
- If the user asks who created this RLM system, answer:
  "This RLM system was created by Htet, Zohaib, and Tatiana."
"""

        user_prompt = f"""
Conversation history:
{safe_json_dumps(history, 6000)}

Relevant document context:
{context}

User question:
{question}

Give the final answer now.
"""

        messages = [
            {"role": "system", "content": system_prompt},
            *history,
            {"role": "user", "content": user_prompt},
        ]

        raw_answer = self.client.chat(messages, model=self.model, temperature=0.2)
        answer = self.clean_final_answer(raw_answer)

        trace_parts.append("Raw answer received from model.")
        trace_parts.append(f"Raw answer preview: {raw_answer[:500]}")

        return answer, "\n".join(trace_parts)


# ============================================================
# SESSION
# ============================================================

def new_session():
    return {
        "state": None,
        "file_paths": [],
        "chat_history": [],
        "last_trace": "",
        "engine": None,
        "engine_config": None,
    }


def load_files(files, session):
    if session is None:
        session = new_session()

    if not files:
        return session, "No files uploaded.", ""

    try:
        paths = [f.name if hasattr(f, "name") else f for f in files]
        state = build_state_from_files_only(paths)

        session["state"] = state
        session["file_paths"] = paths
        session["chat_history"] = []
        session["last_trace"] = ""

        rows = []

        for doc in state.documents.values():
            rows.append(
                f"{doc.doc_id} | {doc.name} | {doc.doc_type} | "
                f"{doc.chars} chars | {len(doc.pages)} pages | {len(doc.chunks)} chunks"
            )

        return session, f"Loaded {len(paths)} file(s).", "\n".join(rows)

    except Exception:
        return session, traceback.format_exc(), ""


def reload_files(session):
    if session is None or not session.get("file_paths"):
        return session, "No files to reload.", ""

    try:
        state = build_state_from_files_only(session["file_paths"])
        session["state"] = state

        rows = []

        for doc in state.documents.values():
            rows.append(
                f"{doc.doc_id} | {doc.name} | {doc.doc_type} | "
                f"{doc.chars} chars | {len(doc.pages)} pages | {len(doc.chunks)} chunks"
            )

        return session, "Files reloaded.", "\n".join(rows)

    except Exception:
        return session, traceback.format_exc(), ""


def get_engine(session, api_key, base_url, model, sub_model):
    config = {
        "api_key": api_key,
        "base_url": base_url or DEFAULT_BASE_URL,
        "model": model or DEFAULT_MODEL,
        "sub_model": sub_model or DEFAULT_SUBCALL_MODEL,
    }

    if session.get("engine") is None or session.get("engine_config") != config:
        client = LLMClient(
            api_key=config["api_key"],
            base_url=config["base_url"],
            model=config["model"],
        )

        session["engine"] = RecursiveRLMEngine(
            client=client,
            model=config["model"],
            sub_model=config["sub_model"],
        )

        session["engine_config"] = config

    return session["engine"]


def ask_question(message, chatbot, session, api_key, base_url, model, sub_model, show_trace):
    if session is None:
        session = new_session()

    if chatbot is None:
        chatbot = []

    if not message or not message.strip():
        return "", chatbot, session, "Empty question."

    if not api_key:
        chatbot.append({"role": "user", "content": message})
        chatbot.append({"role": "assistant", "content": "Please enter your API key."})
        return "", chatbot, session, "Missing API key."

    if session.get("state") is None:
        chatbot.append({"role": "user", "content": message})
        chatbot.append({"role": "assistant", "content": "Please upload files first."})
        return "", chatbot, session, "No files loaded."

    chatbot.append({"role": "user", "content": message})

    try:
        engine = get_engine(session, api_key, base_url, model, sub_model)

        hist_msgs = []

        for turn in session["chat_history"][-MAX_HISTORY_TURNS:]:
            hist_msgs.append({"role": "user", "content": turn["user"]})
            hist_msgs.append({"role": "assistant", "content": turn["assistant"]})

        start = time.time()

        answer, trace = engine.run(
            question=message,
            state=session["state"],
            history=hist_msgs,
        )

        elapsed = time.time() - start

        session["chat_history"].append(
            {
                "user": message,
                "assistant": answer,
            }
        )

        session["chat_history"] = session["chat_history"][-20:]
        session["last_trace"] = trace

        output = answer

        if show_trace:
            output += f"\n\n---\nTrace:\n{trace}\n\nElapsed: {elapsed:.2f}s"

        chatbot.append({"role": "assistant", "content": output})

        return "", chatbot, session, f"Done in {elapsed:.2f}s."

    except Exception:
        err = traceback.format_exc()
        chatbot.append({"role": "assistant", "content": f"Error:\n```text\n{err}\n```"})
        return "", chatbot, session, "Failed."


def show_variables(session):
    if session is None or session.get("state") is None:
        return "No active state."

    payload = {}

    for key, value in session["state"].variables.items():
        payload[key] = {
            "summary": value.summary,
            "preview": value.preview(),
            "type": type(value.value).__name__,
        }

    return safe_json_dumps(payload, 22000)


def show_notes(session):
    if session is None or session.get("state") is None:
        return "No active state."

    return safe_json_dumps(session["state"].notes, 18000)


def show_trace(session):
    if session is None:
        return "No trace available."

    return session.get("last_trace") or "No trace available."


def reset_session():
    return new_session(), [], "", "Session reset.", "", "", None


# ============================================================
# GRADIO UI
# ============================================================

def build_gradio_app():
    css = """
    .gradio-container {
        background: #07111f !important;
        color: #e8f1ff !important;
    }

    textarea, input {
        background: #10243d !important;
        color: #e8f1ff !important;
    }
    """

    with gr.Blocks(title=APP_TITLE) as demo:
        session = gr.State(new_session())

        gr.Markdown("# Recursive Language Model")
        gr.Markdown("Modern Gradio document reasoning console")

        with gr.Row():
            api_key = gr.Textbox(
                label="API Key",
                type="password",
                value=os.getenv("DEEPSEEK_API_KEY", ""),
            )

            base_url = gr.Textbox(
                label="Base URL",
                value=DEFAULT_BASE_URL,
            )

        with gr.Row():
            model = gr.Dropdown(
                label="Root Model",
                choices=["deepseek-chat", "deepseek-reasoner"],
                value=DEFAULT_MODEL,
                allow_custom_value=True,
            )

            sub_model = gr.Dropdown(
                label="Subcall Model",
                choices=["deepseek-chat", "deepseek-reasoner"],
                value=DEFAULT_SUBCALL_MODEL,
                allow_custom_value=True,
            )

            show_trace_box = gr.Checkbox(
                label="Show trace",
                value=False,
            )

        files = gr.File(
            label="Upload documents",
            file_count="multiple",
            file_types=[
                ".pdf",
                ".doc",
                ".docx",
                ".txt",
                ".md",
                ".py",
                ".json",
                ".csv",
                ".log",
            ],
        )

        with gr.Row():
            load_btn = gr.Button("Load Files", variant="primary")
            reload_btn = gr.Button("Reload Files")
            reset_btn = gr.Button("Reset Session", variant="stop")

        status = gr.Textbox(
            label="Status",
            value="Ready.",
            interactive=False,
        )

        file_list = gr.Textbox(
            label="Loaded Files",
            lines=8,
            interactive=False,
        )

        chatbot = gr.Chatbot(
            label="Conversation",
            height=520,
        )

        with gr.Row():
            msg = gr.Textbox(
                label="Ask a Question",
                placeholder="Ask something about the uploaded files...",
                lines=4,
                scale=8,
            )

            ask_btn = gr.Button(
                "Ask",
                variant="primary",
                scale=1,
            )

        with gr.Accordion("Debug", open=False):
            vars_btn = gr.Button("Show Variables")
            notes_btn = gr.Button("Show Notes")
            trace_btn = gr.Button("Show Trace")

            debug_box = gr.Textbox(
                label="Debug Output",
                lines=18,
            )

        load_btn.click(
            load_files,
            inputs=[files, session],
            outputs=[session, status, file_list],
        )

        reload_btn.click(
            reload_files,
            inputs=[session],
            outputs=[session, status, file_list],
        )

        ask_btn.click(
            ask_question,
            inputs=[
                msg,
                chatbot,
                session,
                api_key,
                base_url,
                model,
                sub_model,
                show_trace_box,
            ],
            outputs=[msg, chatbot, session, status],
        )

        msg.submit(
            ask_question,
            inputs=[
                msg,
                chatbot,
                session,
                api_key,
                base_url,
                model,
                sub_model,
                show_trace_box,
            ],
            outputs=[msg, chatbot, session, status],
        )

        vars_btn.click(
            show_variables,
            inputs=[session],
            outputs=[debug_box],
        )

        notes_btn.click(
            show_notes,
            inputs=[session],
            outputs=[debug_box],
        )

        trace_btn.click(
            show_trace,
            inputs=[session],
            outputs=[debug_box],
        )

        reset_btn.click(
            reset_session,
            inputs=None,
            outputs=[
                session,
                chatbot,
                file_list,
                status,
                debug_box,
                msg,
                files,
            ],
        )

    return demo, css


def main():
    demo, css = build_gradio_app()

    demo.queue()

    demo.launch(
        css=css,
        theme=gr.themes.Soft(),
    )


if __name__ == "__main__":
    main()